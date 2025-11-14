"""
Document Watermarking Module
Adds watermarks to PDF, HTML, and DOCX files
Consolidated from private_watermark_system tools
"""

import os
import base64
import json
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict
from io import BytesIO


class Watermarker:
    """Apply watermarks to documents"""

    def __init__(self):
        """Initialize watermarker"""
        self.supported_formats = ['.pdf', '.html', '.htm', '.docx', '.md', '.markdown']

    def get_logo_base64(self, logo_path: Path) -> str:
        """Convert logo image to base64 string"""
        try:
            with open(logo_path, 'rb') as f:
                logo_data = f.read()
                base64_data = base64.b64encode(logo_data).decode('utf-8')

                # Determine MIME type
                ext = logo_path.suffix.lower()
                mime_types = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.gif': 'image/gif',
                    '.bmp': 'image/bmp'
                }
                mime_type = mime_types.get(ext, 'image/png')

                return f"data:{mime_type};base64,{base64_data}"
        except Exception as e:
            print(f"Error encoding logo: {e}")
            return ""

    def generate_watermark_css(self, watermark_text: str, logo_base64: str,
                              opacity: float, position: str) -> str:
        """Generate CSS for HTML watermark"""
        position_styles = {
            'center': 'top: 50%; left: 50%; transform: translate(-50%, -50%);',
            'top-left': 'top: 20px; left: 20px;',
            'top-right': 'top: 20px; right: 20px;',
            'bottom-left': 'bottom: 20px; left: 20px;',
            'bottom-right': 'bottom: 20px; right: 20px;'
        }

        pos_style = position_styles.get(position, position_styles['center'])

        css = f"""
        <style>
        .watermark-overlay {{
            position: fixed;
            {pos_style}
            opacity: {opacity};
            z-index: 9999;
            pointer-events: none;
            text-align: center;
        }}
        .watermark-overlay img {{
            max-width: 150px;
            max-height: 150px;
            display: block;
            margin: 0 auto 10px auto;
        }}
        .watermark-overlay .watermark-text {{
            font-size: 14px;
            color: #333;
            font-family: Arial, sans-serif;
        }}
        @media print {{
            .watermark-overlay {{
                position: absolute;
            }}
        }}
        </style>
        """
        return css

    def watermark_html(self, input_file: Path, watermark_text: str,
                      logo_path: Optional[Path], opacity: float,
                      position: str, output_dir: Path) -> Path:
        """Add watermark to HTML file"""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        # Read HTML content
        with open(input_file, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        # Generate logo base64 if provided
        logo_base64 = ""
        if logo_path and logo_path.exists():
            logo_base64 = self.get_logo_base64(logo_path)

        # Generate watermark CSS
        watermark_css = self.generate_watermark_css(watermark_text, logo_base64, opacity, position)

        # Create watermark HTML
        watermark_html = '<div class="watermark-overlay">'
        if logo_base64:
            watermark_html += f'<img src="{logo_base64}" alt="Watermark Logo">'
        if watermark_text:
            watermark_html += f'<div class="watermark-text">{watermark_text}</div>'
        watermark_html += '</div>'

        # Insert watermark before closing </head> tag or at beginning of <body>
        if '</head>' in html_content:
            html_content = html_content.replace('</head>', f'{watermark_css}</head>')
        elif '<body>' in html_content:
            html_content = html_content.replace('<body>', f'<body>{watermark_css}')
        else:
            # No proper HTML structure, add at beginning
            html_content = watermark_css + html_content

        # Insert watermark div after <body> tag
        if '<body>' in html_content:
            html_content = html_content.replace('<body>', f'<body>\n{watermark_html}\n', 1)
        else:
            html_content = watermark_html + '\n' + html_content

        # Save watermarked HTML
        output_filename = f"{input_file.stem}_watermarked{input_file.suffix}"
        output_path = output_dir / output_filename

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        # Save metadata
        self._save_metadata(output_path, watermark_text, opacity, position)

        return output_path

    def watermark_pdf(self, input_file: Path, watermark_text: str,
                     logo_path: Optional[Path], opacity: float,
                     position: str, output_dir: Path) -> Path:
        """Add watermark to PDF file"""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            from PyPDF2 import PdfReader, PdfWriter, Transformation
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import Color
            from reportlab.lib.utils import ImageReader
        except ImportError:
            raise Exception("PyPDF2 and reportlab are required for PDF watermarking. Install with: pip install PyPDF2 reportlab")

        print(f"[PDF WATERMARK] Reading PDF: {input_file}")

        # Read input PDF
        reader = PdfReader(str(input_file))
        writer = PdfWriter()

        print(f"[PDF WATERMARK] PDF has {len(reader.pages)} pages")

        if not watermark_text:
            raise Exception("Watermark text is required for PDF watermarking")

        # Process each page
        for page_num, page in enumerate(reader.pages):
            print(f"[PDF WATERMARK] Processing page {page_num + 1}/{len(reader.pages)}")

            # Get page dimensions
            page_box = page.mediabox
            page_width = float(page_box.width)
            page_height = float(page_box.height)

            print(f"[PDF WATERMARK] Page size: {page_width} x {page_height}")

            # Create watermark for this page with correct size
            packet = BytesIO()
            can = canvas.Canvas(packet, pagesize=(page_width, page_height))

            # Set opacity and color
            can.setFillColor(Color(0.5, 0.5, 0.5, alpha=opacity))  # Gray color

            # Use larger font for better visibility
            font_size = 14
            can.setFont("Helvetica-Bold", font_size)

            # Get text width if watermark text exists
            text_width = 0
            if watermark_text:
                text_width = can.stringWidth(watermark_text, "Helvetica-Bold", font_size)

            # Logo dimensions
            logo_size = 100

            # Calculate total width/height needed for watermark (text + logo)
            total_width = max(text_width, logo_size)
            total_height = logo_size + 20 + font_size  # logo + spacing + text

            # Calculate base position with proper alignment
            margin = 50
            positions = {
                'center': (page_width / 2 - total_width / 2, page_height / 2),
                'top-left': (margin, page_height - margin - total_height),
                'top-right': (page_width - margin - total_width, page_height - margin - total_height),
                'bottom-left': (margin, margin),
                'bottom-right': (page_width - margin - total_width, margin)
            }

            base_x, base_y = positions.get(position, positions['center'])

            print(f"[PDF WATERMARK] Position: {position} = ({base_x}, {base_y}), text_width={text_width}")

            # Draw logo if provided
            logo_height = 0
            if logo_path and logo_path.exists():
                try:
                    print(f"[PDF WATERMARK] Adding logo: {logo_path}")

                    # Use PIL to create a transparent version of the logo
                    from PIL import Image

                    # Load and resize logo
                    logo_img = Image.open(str(logo_path))

                    # Convert to RGBA if not already
                    if logo_img.mode != 'RGBA':
                        logo_img = logo_img.convert('RGBA')

                    # Set logo size
                    logo_width = 100
                    logo_height = 100
                    logo_img = logo_img.resize((logo_width, logo_height), Image.Resampling.LANCZOS)

                    # Apply opacity to the image
                    # Create a new image with adjusted alpha channel
                    alpha = logo_img.split()[3]  # Get alpha channel
                    # Multiply alpha by opacity
                    alpha = alpha.point(lambda p: int(p * opacity))
                    logo_img.putalpha(alpha)

                    # Save to BytesIO for reportlab
                    logo_io = BytesIO()
                    logo_img.save(logo_io, format='PNG')
                    logo_io.seek(0)

                    # Adjust position for logo
                    logo_x = base_x - logo_width / 2 if position == 'center' else base_x
                    logo_y = base_y + 20  # Place logo above text

                    # Draw logo (opacity already applied via PIL)
                    # Use ImageReader to handle BytesIO
                    img_reader = ImageReader(logo_io)
                    can.drawImage(img_reader, logo_x, logo_y,
                                width=logo_width, height=logo_height,
                                preserveAspectRatio=True, mask='auto')

                    print(f"[PDF WATERMARK] Logo added at ({logo_x}, {logo_y}) with opacity {opacity}")
                    logo_height += 20  # Add spacing
                except Exception as e:
                    print(f"[PDF WATERMARK] Logo error: {e}")
                    import traceback
                    traceback.print_exc()

            # Draw text watermark
            if watermark_text:
                text_x = base_x - len(watermark_text) * 3 if position == 'center' else base_x

                # Adjust text position based on logo placement
                if logo_height > 0:
                    # Logo exists - place text below logo
                    if position in ['bottom-left', 'bottom-right']:
                        # For bottom positions, keep text above page boundary
                        text_y = max(base_y - logo_height, 20)
                    else:
                        text_y = base_y - logo_height
                else:
                    # No logo - place text at base position
                    text_y = base_y

                print(f"[PDF WATERMARK] Text: '{watermark_text}' at ({text_x}, {text_y}) with opacity {opacity}")
                can.drawString(text_x, text_y, watermark_text)
            can.save()

            # Move to beginning of BytesIO buffer
            packet.seek(0)

            # Read watermark PDF
            watermark_pdf = PdfReader(packet)
            watermark_page = watermark_pdf.pages[0]

            # Merge watermark with page
            page.merge_page(watermark_page)

            # Add page to writer
            writer.add_page(page)

        # Save watermarked PDF
        output_filename = f"{input_file.stem}_watermarked{input_file.suffix}"
        output_path = output_dir / output_filename

        print(f"[PDF WATERMARK] Saving to: {output_path}")

        with open(output_path, 'wb') as f:
            writer.write(f)

        print(f"[PDF WATERMARK] File saved: {output_path.exists()}")
        print(f"[PDF WATERMARK] File size: {output_path.stat().st_size} bytes")

        # Save metadata
        self._save_metadata(output_path, watermark_text, opacity, position)

        return output_path

    def watermark_docx(self, input_file: Path, watermark_text: str,
                      logo_path: Optional[Path], opacity: float,
                      position: str, output_dir: Path) -> Path:
        """Add watermark to DOCX file"""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        print(f"[DOCX WATERMARK] Processing: {input_file}")

        output_filename = f"{input_file.stem}_watermarked{input_file.suffix}"
        output_path = output_dir / output_filename

        # Try to use python-docx for proper watermarking
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            print(f"[DOCX WATERMARK] Using python-docx library")

            # Load document
            doc = Document(str(input_file))

            # Add watermark to header of first section
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

            # Set header text
            header_para.text = watermark_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Style the watermark text
            for run in header_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
                run.font.italic = True

            # Add watermark to footer as well
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = watermark_text
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in footer_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.italic = True

            # Save document
            doc.save(str(output_path))

            print(f"[DOCX WATERMARK] Watermark added to header and footer")

        except ImportError:
            print(f"[DOCX WATERMARK] python-docx not available, using fallback method")

            # Fallback: Convert to PDF and watermark the PDF instead
            import shutil
            shutil.copy2(input_file, output_path)

            print(f"[DOCX WATERMARK] WARNING: python-docx not installed")
            print(f"[DOCX WATERMARK] File copied without watermark")
            print(f"[DOCX WATERMARK] Install with: pip install python-docx")

        except Exception as e:
            print(f"[DOCX WATERMARK] Error: {e}")
            # Fallback to simple copy
            import shutil
            shutil.copy2(input_file, output_path)

        # Create HTML companion
        html_companion = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Watermark Information - {input_file.name}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; }}
        .info {{ background: #f0f0f0; padding: 20px; border-radius: 8px; }}
        .watermark {{ opacity: {opacity}; text-align: center; margin: 20px 0; }}
    </style>
</head>
<body>
    <h1>Document Watermark Information</h1>
    <div class="info">
        <p><strong>Original File:</strong> {input_file.name}</p>
        <p><strong>Watermarked:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><strong>Watermark Text:</strong> {watermark_text}</p>
        <p><strong>Opacity:</strong> {opacity * 100}%</p>
        <p><strong>Position:</strong> {position}</p>
    </div>
    <div class="watermark">
        <p>{watermark_text}</p>
    </div>
</body>
</html>"""

        html_path = output_dir / f"{input_file.stem}_watermark_info.html"
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_companion)

        # Save metadata
        self._save_metadata(output_path, watermark_text, opacity, position)

        return output_path

    def watermark_md(self, input_file: Path, watermark_text: str,
                    logo_path: Optional[Path], opacity: float,
                    position: str, output_dir: Path) -> Path:
        """Add watermark to Markdown file as metadata and comments"""
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        print(f"[MD WATERMARK] Processing: {input_file}")

        # Read markdown content
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Create watermark metadata in YAML frontmatter
        watermark_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        copyright_year = datetime.now().year

        # Check if file already has YAML frontmatter
        has_frontmatter = md_content.strip().startswith('---')

        watermark_yaml = f"""---
watermark: "{watermark_text}"
watermarked_date: "{watermark_date}"
copyright: "© {copyright_year} {watermark_text}. All Rights Reserved."
document_id: "{input_file.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
opacity: {opacity}
position: "{position}"
---

"""

        # Add HTML comment watermark
        html_watermark = f"""<!--
WATERMARKED DOCUMENT
==================
Watermark: {watermark_text}
Date: {watermark_date}
Copyright: © {copyright_year} {watermark_text}. All Rights Reserved.
Position: {position}
Opacity: {opacity}

⚠️ NOTICE: This document is watermarked and protected.
Unauthorized distribution or modification is prohibited.
-->

"""

        # If logo provided, embed as base64 in HTML comment
        if logo_path and logo_path.exists():
            logo_base64 = self.get_logo_base64(logo_path)
            if logo_base64:
                html_watermark += f"""<!-- Watermark Logo (Base64):
![Watermark Logo]({logo_base64})
-->

"""

        # Build watermarked content
        if has_frontmatter:
            # Insert watermark data into existing frontmatter
            # Find the end of existing frontmatter
            parts = md_content.split('---', 2)
            if len(parts) >= 3:
                # parts[0] is empty, parts[1] is existing frontmatter, parts[2] is content
                existing_frontmatter = parts[1]
                rest_content = parts[2]

                watermarked_content = f"""---
{existing_frontmatter.strip()}
# Watermark Information
watermark: "{watermark_text}"
watermarked_date: "{watermark_date}"
copyright: "© {copyright_year} {watermark_text}. All Rights Reserved."
---

{html_watermark}{rest_content}"""
            else:
                # Malformed frontmatter, add new one
                watermarked_content = watermark_yaml + html_watermark + md_content
        else:
            # No existing frontmatter, add new one
            watermarked_content = watermark_yaml + html_watermark + md_content

        # Add footer watermark
        footer_watermark = f"""

---

**Document Watermark Information**

- Watermarked: {watermark_date}
- {watermark_text}
- © {copyright_year} All Rights Reserved

*This document is protected and watermarked. Unauthorized distribution is prohibited.*

<!-- End of Watermark -->
"""

        watermarked_content += footer_watermark

        # Save watermarked markdown
        output_filename = f"{input_file.stem}_watermarked{input_file.suffix}"
        output_path = output_dir / output_filename

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(watermarked_content)

        print(f"[MD WATERMARK] Saved: {output_path}")

        # Save metadata
        self._save_metadata(output_path, watermark_text, opacity, position)

        return output_path

    def _save_metadata(self, output_path: Path, watermark_text: str,
                      opacity: float, position: str):
        """Save watermark metadata to JSON file"""
        metadata = {
            'file': str(output_path.name),
            'watermark_text': watermark_text,
            'opacity': opacity,
            'position': position,
            'timestamp': datetime.now().isoformat(),
            'copyright': f"© {datetime.now().year} Watermarked Document. All Rights Reserved."
        }

        metadata_path = output_path.parent / f"{output_path.stem}.watermark.json"
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)

    def apply_watermark(self, input_file: Path, watermark_text: str,
                       logo_path: Optional[Path], opacity: float,
                       position: str, output_dir: Path) -> Path:
        """
        Apply watermark to document

        Args:
            input_file: Path to input document
            watermark_text: Text to use as watermark
            logo_path: Optional path to logo image
            opacity: Watermark opacity (0.0 to 1.0)
            position: Watermark position (center, top-left, etc.)
            output_dir: Output directory

        Returns:
            Path to watermarked document
        """
        input_file = Path(input_file)
        output_dir = Path(output_dir)

        if not input_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_file}")

        ext = input_file.suffix.lower()

        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}. Supported formats: {self.supported_formats}")

        # Ensure opacity is in valid range
        opacity = max(0.0, min(1.0, opacity))

        # Apply watermark based on file type
        if ext in ['.html', '.htm']:
            return self.watermark_html(input_file, watermark_text, logo_path, opacity, position, output_dir)
        elif ext == '.pdf':
            return self.watermark_pdf(input_file, watermark_text, logo_path, opacity, position, output_dir)
        elif ext == '.docx':
            return self.watermark_docx(input_file, watermark_text, logo_path, opacity, position, output_dir)
        elif ext in ['.md', '.markdown']:
            return self.watermark_md(input_file, watermark_text, logo_path, opacity, position, output_dir)
        else:
            raise ValueError(f"Unsupported format: {ext}")


if __name__ == '__main__':
    # Test watermarker
    watermarker = Watermarker()
    print("Watermarker initialized")
    print(f"Supported formats: {watermarker.supported_formats}")
