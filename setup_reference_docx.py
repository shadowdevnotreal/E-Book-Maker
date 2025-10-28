#!/usr/bin/env python3
"""
Setup Reference DOCX with Page Numbering
Programmatically adds page numbers to Pandoc reference document
"""

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_page_number_to_footer(doc):
    """Add page numbers to all footers in the document"""
    for section in doc.sections:
        # Access the footer
        footer = section.footer

        # Clear existing footer content
        for element in footer.paragraphs:
            element.clear()

        # Create a new paragraph for page numbers
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the page number field
        run = paragraph.add_run()

        # Create the PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        # Add the field to the run
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Set font properties
        run.font.size = Pt(11)
        run.font.name = 'Calibri'


def setup_reference_document():
    """Set up the reference DOCX with page numbering"""
    reference_path = Path(__file__).parent / 'config' / 'templates' / 'reference-with-pagenumbers.docx'

    if not reference_path.exists():
        print(f"Error: Reference document not found at {reference_path}")
        print("Please run: pandoc --print-default-data-file reference.docx > config/templates/reference-with-pagenumbers.docx")
        return False

    print(f"Setting up reference document: {reference_path}")

    try:
        # Load the document
        doc = Document(str(reference_path))

        # Add page numbers to footer
        add_page_number_to_footer(doc)

        # Save the document
        doc.save(str(reference_path))

        print(f"✓ Successfully added page numbers to {reference_path.name}")
        print(f"  Page numbers will appear in the footer (centered)")
        print(f"  This template will be used for all DOCX conversions")

        return True

    except Exception as e:
        print(f"Error setting up reference document: {e}")
        print(f"\nAlternative: Open the file in Microsoft Word and:")
        print(f"  1. Go to Insert > Page Number > Bottom of Page > Plain Number 2 (centered)")
        print(f"  2. Save and close the document")
        return False


if __name__ == '__main__':
    import sys

    # Check if python-docx is installed
    try:
        import docx
    except ImportError:
        print("Error: python-docx is not installed")
        print("Install it with: pip install python-docx")
        sys.exit(1)

    success = setup_reference_document()
    sys.exit(0 if success else 1)
