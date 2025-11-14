#!/usr/bin/env python3
"""
MASTER TEST SCRIPT - E-Book Maker
Tests ALL features and outputs to desktop (not repo)
"""

import sys
import json
import time
from pathlib import Path
from datetime import datetime

# Add modules to path
BASE_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(BASE_DIR / 'modules'))

from ai.groq_assistant import GroqAssistant
from conversion.converter import EBookConverter
from covers.cover_generator import CoverGenerator
from watermarking.watermarker import Watermarker

# Output to desktop (NOT in repo)
# Use Windows desktop path for WSL compatibility
TEST_DIR = Path("/mnt/c/Users/Mishka/Desktop") / "E-Book-Maker-Tests" / f"full_test_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
TEST_DIR.mkdir(parents=True, exist_ok=True)

# Create subdirectories
(TEST_DIR / "covers").mkdir(exist_ok=True)
(TEST_DIR / "ebooks").mkdir(exist_ok=True)
(TEST_DIR / "watermarked").mkdir(exist_ok=True)
(TEST_DIR / "ai_outputs").mkdir(exist_ok=True)

# Test results
results = {
    "timestamp": datetime.now().isoformat(),
    "test_categories": {}
}


def log_test(category, test_name, success, details="", output_files=None):
    """Log test result"""
    if category not in results["test_categories"]:
        results["test_categories"][category] = []

    result = {
        "test": test_name,
        "success": success,
        "details": details,
        "output_files": output_files or []
    }
    results["test_categories"][category].append(result)

    status = "✓" if success else "✗"
    print(f"  {status} {test_name}")
    if details:
        print(f"    → {details}")
    if output_files:
        for f in output_files:
            print(f"    💾 {f}")


def create_test_content():
    """Create test content files"""
    print("\n📝 Creating test content files...")

    # Create sample markdown
    sample_md = TEST_DIR / "sample.md"
    sample_md.write_text("""# Test Book: E-Book Maker Features

## Chapter 1: Introduction

This is a comprehensive test of the E-Book Maker system. It includes multiple chapters, formatting, and various text elements.

### Key Features Tested:
- Document conversion
- Cover generation
- Watermarking
- AI assistance
- Smart text optimization

## Chapter 2: Advanced Features

The E-Book Maker includes state-of-the-art features:

1. **AI Integration** - Groq-powered intelligence
2. **Smart Text** - Automatic readability optimization
3. **KDP Compliance** - Amazon-ready outputs

### Technical Specifications

The system handles multiple formats including:
- EPUB (e-reader format)
- PDF (print-ready, 300 DPI)
- DOCX (Microsoft Word)
- HTML (web format)
- Markdown (plain text)

## Chapter 3: Conclusion

This test demonstrates the complete capabilities of the E-Book Maker system.

> "The best e-book creation tool for Amazon KDP publishing." - Test User

### Call to Action

Ready to create your own e-books? Let's get started!

---

**About the Author**

This is a test document generated for comprehensive system testing.
""")

    print(f"  ✓ Created: {sample_md}")
    return sample_md


def test_document_conversion(sample_file):
    """Test all document conversion features"""
    print("\n" + "=" * 80)
    print("📄 TESTING: Document Conversion")
    print("=" * 80)

    try:
        converter = EBookConverter()

        # Test each output format
        formats_to_test = ['epub', 'pdf', 'html', 'docx', 'md']

        for fmt in formats_to_test:
            try:
                print(f"\n  Converting to {fmt.upper()}...")
                output_files = converter.convert(
                    [sample_file],
                    [fmt],
                    title="Test E-Book: Feature Validation",
                    author="E-Book Maker Test Suite",
                    subtitle="Comprehensive System Testing",
                    output_dir=TEST_DIR / "ebooks"
                )

                output_paths = []
                for file in output_files:
                    if Path(file).exists():
                        output_paths.append(str(Path(file).relative_to(TEST_DIR)))

                log_test(
                    "Document Conversion",
                    f"Convert to {fmt.upper()}",
                    len(output_paths) > 0,
                    f"Generated {len(output_paths)} file(s)",
                    output_paths
                )

            except Exception as e:
                log_test("Document Conversion", f"Convert to {fmt.upper()}", False, str(e))

        return True

    except Exception as e:
        print(f"  ✗ Conversion test failed: {e}")
        return False


def test_cover_generation():
    """Test all cover generation features"""
    print("\n" + "=" * 80)
    print("🎨 TESTING: Cover Generation")
    print("=" * 80)

    try:
        generator = CoverGenerator()

        # Test 1: E-Book Cover (standard gradient)
        print("\n  Generating e-book cover (gradient)...")
        try:
            ebook_cover = generator.create_cover(
                cover_type='ebook',
                title='E-Book Maker',
                subtitle='Professional Publishing Suite',
                author='Test Suite',
                style='gradient',
                colors={'primary': '#667eea', 'secondary': '#764ba2'},
                output_dir=TEST_DIR / "covers"
            )
            log_test(
                "Cover Generation",
                "E-Book Cover (Gradient)",
                ebook_cover.exists(),
                f"1600x2560 px, 300 DPI",
                [str(ebook_cover.relative_to(TEST_DIR))]
            )
        except Exception as e:
            log_test("Cover Generation", "E-Book Cover (Gradient)", False, str(e))

        # Test 2: E-Book Cover with long title (text wrapping)
        print("\n  Generating e-book cover with long title (text wrapping test)...")
        try:
            long_title_cover = generator.create_cover(
                cover_type='ebook',
                title='Complete Professional E-Book Creation System',
                subtitle='Featuring AI-Powered Content Generation and Smart Text Optimization',
                author='Comprehensive Test Suite',
                style='solid',
                colors={'primary': '#FF5733', 'secondary': '#C70039'},
                output_dir=TEST_DIR / "covers"
            )
            log_test(
                "Cover Generation",
                "E-Book Cover (Text Wrapping)",
                long_title_cover.exists(),
                "Long title with automatic wrapping",
                [str(long_title_cover.relative_to(TEST_DIR))]
            )
        except Exception as e:
            log_test("Cover Generation", "E-Book Cover (Text Wrapping)", False, str(e))

        # Test 3: E-Book Cover with custom background (test background image support)
        print("\n  Generating e-book cover with solid color background...")
        try:
            from PIL import Image
            # Create a test background
            test_bg = Image.new('RGB', (1600, 2560), (255, 200, 100))
            bg_path = TEST_DIR / "covers" / "test_background.jpg"
            test_bg.save(bg_path, 'JPEG', quality=95)

            bg_cover = generator.create_cover(
                cover_type='ebook',
                title='Background Image Test',
                subtitle='Smart Text Color Optimization',
                author='Test Suite',
                style='gradient',
                colors={'primary': '#667eea', 'secondary': '#764ba2'},
                output_dir=TEST_DIR / "covers",
                background_image=bg_path
            )
            log_test(
                "Cover Generation",
                "E-Book Cover (Custom Background)",
                bg_cover.exists(),
                "With background image and smart text color",
                [str(bg_cover.relative_to(TEST_DIR))]
            )
        except Exception as e:
            log_test("Cover Generation", "E-Book Cover (Custom Background)", False, str(e))

        # Test 4: Minimalist style
        print("\n  Generating minimalist cover...")
        try:
            minimal_cover = generator.create_cover(
                cover_type='ebook',
                title='Minimalist Design',
                subtitle='Clean and Professional',
                author='Test Suite',
                style='minimalist',
                colors={'primary': '#000000', 'secondary': '#FFFFFF'},
                output_dir=TEST_DIR / "covers"
            )
            log_test(
                "Cover Generation",
                "E-Book Cover (Minimalist)",
                minimal_cover.exists(),
                "Clean white background with black text",
                [str(minimal_cover.relative_to(TEST_DIR))]
            )
        except Exception as e:
            log_test("Cover Generation", "E-Book Cover (Minimalist)", False, str(e))

        return True

    except Exception as e:
        print(f"  ✗ Cover generation test failed: {e}")
        return False


def test_watermarking():
    """Test watermarking features"""
    print("\n" + "=" * 80)
    print("💧 TESTING: Watermarking")
    print("=" * 80)

    try:
        watermarker = Watermarker()

        # Create test files to watermark
        test_files = []

        # Create test PDF
        print("\n  Creating test PDF...")
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        pdf_path = TEST_DIR / "test_document.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test Document for Watermarking")
        c.drawString(100, 700, "This is a sample PDF document.")
        c.showPage()
        c.save()
        test_files.append(pdf_path)

        # Create test HTML
        html_path = TEST_DIR / "test_document.html"
        html_path.write_text("""<!DOCTYPE html>
<html>
<head><title>Test Document</title></head>
<body>
    <h1>Test Document for Watermarking</h1>
    <p>This is a sample HTML document.</p>
</body>
</html>""")
        test_files.append(html_path)

        # Create test DOCX
        print("  Creating test DOCX...")
        from docx import Document
        doc = Document()
        doc.add_heading('Test Document for Watermarking', 0)
        doc.add_paragraph('This is a sample Word document.')
        docx_path = TEST_DIR / "test_document.docx"
        doc.save(str(docx_path))
        test_files.append(docx_path)

        # Create test Markdown
        md_path = TEST_DIR / "test_document.md"
        md_path.write_text("""# Test Document for Watermarking

This is a sample Markdown document.

## Features
- Clean formatting
- Easy to read
""")
        test_files.append(md_path)

        # Test watermarking each file type
        for test_file in test_files:
            file_type = test_file.suffix[1:].upper()
            print(f"\n  Watermarking {file_type}...")
            try:
                output_path = watermarker.apply_watermark(
                    input_file=test_file,
                    watermark_text="TEST COPY - E-Book Maker",
                    logo_path=None,
                    opacity=0.3,
                    position='center',
                    output_dir=TEST_DIR / "watermarked"
                )

                log_test(
                    "Watermarking",
                    f"Watermark {file_type}",
                    output_path.exists(),
                    "Text watermark applied",
                    [str(output_path.relative_to(TEST_DIR))]
                )
            except Exception as e:
                log_test("Watermarking", f"Watermark {file_type}", False, str(e))

        return True

    except Exception as e:
        print(f"  ✗ Watermarking test failed: {e}")
        return False


def test_ai_features(api_key):
    """Test all AI features"""
    print("\n" + "=" * 80)
    print("🤖 TESTING: AI Features (Groq Integration)")
    print("=" * 80)

    try:
        ai = GroqAssistant()

        # Initialize AI
        print("\n  Initializing AI with API key...")
        success, message = ai.set_api_key(api_key)
        log_test("AI Features", "API Key Validation", success, message)

        if not ai.is_enabled():
            print("  ✗ AI not enabled, skipping AI tests")
            return False

        ai_output_dir = TEST_DIR / "ai_outputs"

        # Test Cover Design AI
        print("\n  Testing Cover Design AI...")

        title = ai.suggest_cover_title("Machine Learning for Beginners", "Technology")
        log_test("AI Features", "AI Title Suggestion", bool(title), title)
        if title:
            (ai_output_dir / "ai_title.txt").write_text(title)

        if title:
            subtitle = ai.suggest_cover_subtitle(title, "Machine Learning")
            log_test("AI Features", "AI Subtitle Suggestion", bool(subtitle), subtitle)
            if subtitle:
                (ai_output_dir / "ai_subtitle.txt").write_text(subtitle)

        colors = ai.suggest_color_scheme("Technology", "modern")
        log_test("AI Features", "AI Color Scheme", bool(colors), str(colors))
        if colors:
            (ai_output_dir / "ai_colors.json").write_text(json.dumps(colors, indent=2))

        style = ai.suggest_cover_style("Technology", "developers")
        log_test("AI Features", "AI Style Suggestion", bool(style), style)

        # Test Content Generation AI
        print("\n  Testing Content Generation AI...")

        description = ai.generate_book_description(
            "Machine Learning Mastery",
            "From Beginner to Expert",
            "Machine Learning",
            "aspiring data scientists",
            ["Hands-on projects", "Real datasets", "Production deployment"]
        )
        log_test("AI Features", "AI Book Description", bool(description),
                f"{len(description)} chars" if description else "")
        if description:
            (ai_output_dir / "ai_book_description.html").write_text(description)

        bio = ai.generate_author_bio(
            "Dr. Jane Smith",
            "Machine Learning and Data Science",
            "PhD in CS, 15 years industry experience, published researcher"
        )
        log_test("AI Features", "AI Author Bio", bool(bio), bio[:100] if bio else "")
        if bio:
            (ai_output_dir / "ai_author_bio.txt").write_text(bio)

        outline = ai.generate_chapter_outline("Machine Learning Fundamentals", 10)
        log_test("AI Features", "AI Chapter Outline", bool(outline) and len(outline) == 10,
                f"{len(outline)} chapters" if outline else "")
        if outline:
            (ai_output_dir / "ai_chapter_outline.txt").write_text("\n".join(
                [f"{i+1}. {ch}" for i, ch in enumerate(outline)]
            ))

        # Test Text Enhancement AI
        print("\n  Testing Text Enhancement AI...")

        sample_text = "This text have some error and need to be fix for better quality."

        proofread = ai.proofread_text(sample_text)
        log_test("AI Features", "AI Proofreading", bool(proofread), proofread)
        if proofread:
            (ai_output_dir / "ai_proofread.txt").write_text(f"Original: {sample_text}\n\nProofread: {proofread}")

        improved = ai.improve_readability(sample_text, target_grade_level=6)
        log_test("AI Features", "AI Readability", bool(improved), improved)

        expanded = ai.expand_text("Machine learning is important.", target_length=200)
        log_test("AI Features", "AI Text Expansion", bool(expanded),
                f"{len(expanded)} chars" if expanded else "")
        if expanded:
            (ai_output_dir / "ai_expanded_text.txt").write_text(expanded)

        # Test Metadata & Marketing AI
        print("\n  Testing Metadata & Marketing AI...")

        keywords = ai.generate_kdp_keywords("Machine Learning Mastery", "ML education", "Technology")
        log_test("AI Features", "AI KDP Keywords", bool(keywords) and len(keywords) == 7,
                ", ".join(keywords[:3]) + "..." if keywords else "")
        if keywords:
            (ai_output_dir / "ai_kdp_keywords.txt").write_text("\n".join(keywords))

        categories = ai.suggest_kdp_categories("Machine Learning Mastery", "ML", "Technology")
        log_test("AI Features", "AI KDP Categories", bool(categories),
                f"{len(categories)} categories" if categories else "")
        if categories:
            (ai_output_dir / "ai_kdp_categories.txt").write_text("\n".join(categories))

        marketing = ai.generate_marketing_copy(
            "Machine Learning Mastery",
            "Learn ML from scratch to production",
            "Start your ML journey today!"
        )
        log_test("AI Features", "AI Marketing Copy", bool(marketing), marketing[:100] if marketing else "")
        if marketing:
            (ai_output_dir / "ai_marketing.txt").write_text(marketing)

        return True

    except Exception as e:
        print(f"  ✗ AI features test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def generate_report():
    """Generate comprehensive test report"""
    print("\n" + "=" * 80)
    print("📊 GENERATING TEST REPORT")
    print("=" * 80)

    # Count results
    total_tests = 0
    passed_tests = 0

    for category, tests in results["test_categories"].items():
        category_passed = sum(1 for t in tests if t["success"])
        category_total = len(tests)
        total_tests += category_total
        passed_tests += category_passed

        print(f"\n{category}:")
        print(f"  ✓ Passed: {category_passed}/{category_total}")
        print(f"  ✗ Failed: {category_total - category_passed}/{category_total}")

    # Overall summary
    print("\n" + "=" * 80)
    print("OVERALL SUMMARY")
    print("=" * 80)
    print(f"✓ Total Passed: {passed_tests}/{total_tests} ({passed_tests/total_tests*100:.1f}%)")
    print(f"✗ Total Failed: {total_tests - passed_tests}/{total_tests}")

    # Save results
    results_file = TEST_DIR / "test_results.json"
    with open(results_file, 'w') as f:
        json.dump(results, f, indent=2)

    # Create HTML report
    html_report = f"""<!DOCTYPE html>
<html>
<head>
    <title>E-Book Maker - Test Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #333; }}
        .summary {{ background: #f0f0f0; padding: 20px; border-radius: 8px; margin: 20px 0; }}
        .category {{ margin: 20px 0; border-left: 4px solid #667eea; padding-left: 20px; }}
        .success {{ color: #10b981; }}
        .failure {{ color: #ef4444; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #ddd; }}
        th {{ background-color: #667eea; color: white; }}
        .file-list {{ font-size: 12px; color: #666; margin-left: 20px; }}
    </style>
</head>
<body>
    <h1>E-Book Maker - Comprehensive Test Report</h1>
    <div class="summary">
        <h2>Summary</h2>
        <p><strong>Date:</strong> {results['timestamp']}</p>
        <p><strong>Total Tests:</strong> {total_tests}</p>
        <p class="success"><strong>✓ Passed:</strong> {passed_tests} ({passed_tests/total_tests*100:.1f}%)</p>
        <p class="failure"><strong>✗ Failed:</strong> {total_tests - passed_tests}</p>
    </div>
"""

    for category, tests in results["test_categories"].items():
        html_report += f"""
    <div class="category">
        <h2>{category}</h2>
        <table>
            <tr>
                <th>Test</th>
                <th>Status</th>
                <th>Details</th>
            </tr>
"""
        for test in tests:
            status_class = "success" if test["success"] else "failure"
            status_icon = "✓" if test["success"] else "✗"
            html_report += f"""
            <tr>
                <td>{test['test']}</td>
                <td class="{status_class}">{status_icon}</td>
                <td>{test['details']}</td>
            </tr>
"""
            if test['output_files']:
                html_report += f"""
            <tr>
                <td colspan="3" class="file-list">
                    Output files:<br>
                    {"<br>".join(['• ' + f for f in test['output_files']])}
                </td>
            </tr>
"""
        html_report += """
        </table>
    </div>
"""

    html_report += """
</body>
</html>
"""

    html_file = TEST_DIR / "test_report.html"
    html_file.write_text(html_report)

    print(f"\n💾 JSON results: {results_file}")
    print(f"💾 HTML report: {html_file}")
    print(f"\n📁 All test outputs: {TEST_DIR}")

    return passed_tests == total_tests


def main():
    """Run all tests"""
    print("=" * 80)
    print("E-BOOK MAKER - COMPREHENSIVE FEATURE TEST")
    print("=" * 80)
    print(f"Output Directory: {TEST_DIR}")
    print("=" * 80)
    print("\nThis will test ALL features:")
    print("  • Document Conversion (5 formats)")
    print("  • Cover Generation (4 variations)")
    print("  • Watermarking (4 file types)")
    print("  • AI Features (15+ tests)")
    print("  • Smart Text & Optimization")
    print("\nAll outputs will be saved to the desktop test folder.")
    print("=" * 80)

    start_time = time.time()

    # Create test content
    sample_file = create_test_content()

    # Run all tests
    test_document_conversion(sample_file)
    test_cover_generation()
    test_watermarking()
    test_ai_features("")

    # Generate report
    all_passed = generate_report()

    elapsed_time = time.time() - start_time
    print(f"\n⏱️  Total test time: {elapsed_time:.1f} seconds")

    if all_passed:
        print("\n🎉 ALL TESTS PASSED! E-Book Maker is working perfectly!")
    else:
        print("\n⚠️  Some tests failed. Check the report for details.")

    print("\n" + "=" * 80)
    print(f"📁 Open test folder: {TEST_DIR}")
    print("=" * 80)

    return all_passed


if __name__ == "__main__":
    try:
        success = main()
        sys.exit(0 if success else 1)
    except KeyboardInterrupt:
        print("\n\n⚠️  Tests interrupted by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n✗ CRITICAL ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
