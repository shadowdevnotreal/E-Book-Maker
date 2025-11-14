#!/usr/bin/env python3
"""
Comprehensive Test: Watermarking, Text Normalization, and Page Numbering
Tests all advanced document processing features
"""

import sys
from pathlib import Path
from datetime import datetime

# Add modules to path
PROJECT_DIR = Path("/mnt/c/Users/Mishka/Desktop/PROJECTS/E-Book Project/E-Book-Maker")
sys.path.insert(0, str(PROJECT_DIR / 'modules'))

from watermarking.watermarker import Watermarker
from conversion.text_normalizer import TextNormalizer
from conversion.page_numbering import PageNumberingConfig
from conversion.converter import EBookConverter

# Output directory
TEST_DIR = Path("/mnt/c/Users/Mishka/Desktop/E-Book-Maker-Tests") / f"feature_test_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
TEST_DIR.mkdir(parents=True, exist_ok=True)

print("=" * 80)
print("COMPREHENSIVE FEATURE TEST")
print("=" * 80)
print(f"Output: {TEST_DIR}")
print("=" * 80)


def test_watermarking():
    """Test watermarking on all supported formats"""
    print("\n" + "=" * 80)
    print("💧 TEST 1: WATERMARKING")
    print("=" * 80)

    watermarker = Watermarker()
    results = []

    # Create test directory
    watermark_dir = TEST_DIR / "watermarked"
    watermark_dir.mkdir(exist_ok=True)

    # Create test files
    test_files = []

    # 1. Test PDF
    print("\n  Creating test PDF...")
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        pdf_path = TEST_DIR / "test_document.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(100, 750, "Test Document for Watermarking")
        c.drawString(100, 700, "Page 1 - This is a sample PDF document.")
        c.drawString(100, 650, "Testing watermark functionality.")
        c.showPage()

        c.drawString(100, 750, "Page 2 - Second page of test document")
        c.drawString(100, 700, "Watermark should appear on all pages.")
        c.showPage()
        c.save()

        test_files.append(pdf_path)
        print("  ✓ PDF created with 2 pages")
    except Exception as e:
        print(f"  ✗ PDF creation failed: {e}")

    # 2. Test HTML
    print("\n  Creating test HTML...")
    html_path = TEST_DIR / "test_document.html"
    html_content = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Test Document</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; }
        h1 { color: #333; }
        p { line-height: 1.6; }
    </style>
</head>
<body>
    <h1>Test Document for Watermarking</h1>
    <p>This is a sample HTML document created for testing the watermarking feature.</p>
    <p>The watermark should appear as an overlay on this page.</p>
    <h2>Features to Test</h2>
    <ul>
        <li>Text watermark visibility</li>
        <li>Logo watermark (if provided)</li>
        <li>Opacity control</li>
        <li>Position accuracy</li>
    </ul>
</body>
</html>"""
    html_path.write_text(html_content, encoding='utf-8')
    test_files.append(html_path)
    print("  ✓ HTML created")

    # 3. Test DOCX
    print("\n  Creating test DOCX...")
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading('Test Document for Watermarking', 0)
        doc.add_paragraph('This is a sample Word document created for testing.')
        doc.add_paragraph('The watermark should appear in the header and footer.')

        # Add more content
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('Testing watermark on multiple pages.')

        doc.add_page_break()
        doc.add_heading('Section 2', level=1)
        doc.add_paragraph('Page 2 content with watermark.')

        docx_path = TEST_DIR / "test_document.docx"
        doc.save(str(docx_path))
        test_files.append(docx_path)
        print("  ✓ DOCX created with 2 pages")
    except Exception as e:
        print(f"  ✗ DOCX creation failed: {e}")

    # 4. Test Markdown
    print("\n  Creating test Markdown...")
    md_path = TEST_DIR / "test_document.md"
    md_content = """# Test Document for Watermarking

This is a sample Markdown document created for testing the watermarking feature.

## Features

- Text watermark with metadata
- YAML frontmatter
- Copyright information
- Document ID

## Content

The watermark should be embedded in the document metadata and visible in the rendered output.

### Multiple Sections

Testing watermark across multiple sections of content.

---

End of test document.
"""
    md_path.write_text(md_content, encoding='utf-8')
    test_files.append(md_path)
    print("  ✓ Markdown created")

    # Test watermarking each file
    print("\n  Applying watermarks...")
    for test_file in test_files:
        file_type = test_file.suffix[1:].upper()
        print(f"\n  → Testing {file_type} watermark...")

        try:
            output_path = watermarker.apply_watermark(
                input_file=test_file,
                watermark_text="© 2025 E-Book Maker - TEST COPY",
                logo_path=None,
                opacity=0.3,
                position='center',
                output_dir=watermark_dir
            )

            if output_path.exists():
                size = output_path.stat().st_size
                print(f"    ✓ {file_type} watermarked: {output_path.name} ({size:,} bytes)")
                results.append((file_type, True, output_path))
            else:
                print(f"    ✗ {file_type} failed: Output file not created")
                results.append((file_type, False, None))

        except Exception as e:
            print(f"    ✗ {file_type} error: {e}")
            results.append((file_type, False, None))

    # Summary
    passed = sum(1 for _, success, _ in results if success)
    total = len(results)
    print(f"\n  WATERMARKING: {passed}/{total} tests passed")

    return results


def test_text_normalization():
    """Test text normalization features"""
    print("\n" + "=" * 80)
    print("📝 TEST 2: TEXT NORMALIZATION")
    print("=" * 80)

    normalizer = TextNormalizer()
    results = []

    # Create test directory
    normalize_dir = TEST_DIR / "normalized"
    normalize_dir.mkdir(exist_ok=True)

    # Test cases with problematic characters
    test_cases = [
        ("Unicode Quotes", "This is a \u201Ctest\u201D with \u2018smart quotes\u2019 and an em\u2014dash"),
        ("Special Symbols", "Temperature: 25\u00B0C, Copyright \u00A9 2025, 5 \u00D7 3 = 15"),
        ("Math Operators", "x \u2265 5 and y \u2264 10, x \u2260 y, \u03C0 \u2248 3.14"),
        ("Whitespace", "Line\u00A0with\u00A0non-breaking\u00A0spaces and\ttabs"),
        ("List Formatting", "Items:\n- First item\n- Second item\n- Third item"),
        ("Ellipsis", "This is a test\u2026 with ellipsis and other\u2026 characters")
    ]

    print("\n  Testing character normalization...")
    for name, test_text in test_cases:
        print(f"\n  → {name}")
        print(f"    Input:  {repr(test_text)}")

        try:
            normalized = normalizer.normalize_text(test_text)
            print(f"    Output: {repr(normalized)}")

            # Check if problematic characters were replaced
            has_problematic = any(ord(c) > 127 and c not in ['°', '©', '®'] for c in test_text)
            was_normalized = test_text != normalized

            if has_problematic and was_normalized:
                print(f"    ✓ Successfully normalized")
                results.append((name, True))
            elif not has_problematic:
                print(f"    ✓ No normalization needed")
                results.append((name, True))
            else:
                print(f"    ⚠ No changes detected")
                results.append((name, False))

        except Exception as e:
            print(f"    ✗ Error: {e}")
            results.append((name, False))

    # Test file normalization
    print("\n  Testing file normalization...")
    test_md = normalize_dir / "test_input.md"
    test_content = """# Test Document with Special Characters

This document contains \u201Csmart quotes\u201D, em\u2014dashes, and other special characters.

## Mathematical Symbols

- x \u2265 5 (greater than or equal)
- y \u2264 10 (less than or equal)
- x \u2260 y (not equal)
- \u03C0 \u2248 3.14 (approximately)

## Special Symbols

- Temperature: 25\u00B0C
- Copyright \u00A9 2025
- Registered \u00AE trademark

## List with\u00A0spaces

- Item\u00A0one
- Item\u00A0two
- Item\u00A0three

Testing ellipsis\u2026 and other issues.
"""

    test_md.write_text(test_content, encoding='utf-8')

    try:
        output_path = normalizer.normalize_file(test_md)

        if output_path.exists():
            normalized_content = output_path.read_text(encoding='utf-8')
            print(f"    ✓ File normalized: {output_path.name}")
            print(f"    Original size: {len(test_content)} chars")
            print(f"    Normalized size: {len(normalized_content)} chars")
            results.append(("File Normalization", True))
        else:
            print(f"    ✗ Output file not created")
            results.append(("File Normalization", False))

    except Exception as e:
        print(f"    ✗ Error: {e}")
        results.append(("File Normalization", False))

    # Summary
    passed = sum(1 for _, success in results if success)
    total = len(results)
    print(f"\n  TEXT NORMALIZATION: {passed}/{total} tests passed")

    return results


def test_page_numbering():
    """Test page numbering configuration"""
    print("\n" + "=" * 80)
    print("📄 TEST 3: PAGE NUMBERING")
    print("=" * 80)

    results = []

    # Test 1: Default configuration
    print("\n  → Testing default configuration...")
    try:
        config = PageNumberingConfig()
        latex_header = config.generate_latex_header()

        if latex_header:
            print(f"    ✓ LaTeX header generated ({len(latex_header)} chars)")
            print(f"    Headers:\n{latex_header}")
            results.append(("Default Config", True))
        else:
            print(f"    ⚠ Empty LaTeX header")
            results.append(("Default Config", False))
    except Exception as e:
        print(f"    ✗ Error: {e}")
        results.append(("Default Config", False))

    # Test 2: Custom configuration - Footer Right, Roman numerals
    print("\n  → Testing custom configuration (footer-right, roman)...")
    try:
        custom_config = PageNumberingConfig({
            'pdf': {
                'position': 'footer-right',
                'style': 'roman',
                'front_matter': {
                    'enabled': True,
                    'style': 'Roman'
                }
            }
        })

        latex_header = custom_config.generate_latex_header()
        frontmatter_cmd = custom_config.generate_frontmatter_commands()
        mainmatter_cmd = custom_config.generate_mainmatter_commands()

        print(f"    ✓ LaTeX header: {len(latex_header)} chars")
        print(f"    ✓ Front matter: {frontmatter_cmd}")
        print(f"    ✓ Main matter: {mainmatter_cmd}")
        results.append(("Custom Config", True))

    except Exception as e:
        print(f"    ✗ Error: {e}")
        results.append(("Custom Config", False))

    # Test 3: Custom headers and footers
    print("\n  → Testing custom headers/footers...")
    try:
        custom_config = PageNumberingConfig({
            'pdf': {
                'custom_headers': {
                    'enabled': True,
                    'left': 'Chapter Title',
                    'center': '',
                    'right': 'Section'
                },
                'custom_footers': {
                    'enabled': True,
                    'left': 'Author Name',
                    'center': '\\thepage',
                    'right': 'Book Title'
                }
            }
        })

        latex_header = custom_config.generate_latex_header()

        # Check if custom elements are present
        has_header = 'Chapter Title' in latex_header
        has_footer = 'Author Name' in latex_header and 'Book Title' in latex_header

        if has_header and has_footer:
            print(f"    ✓ Custom headers/footers configured")
            results.append(("Custom Headers/Footers", True))
        else:
            print(f"    ⚠ Headers/footers not properly configured")
            results.append(("Custom Headers/Footers", False))

    except Exception as e:
        print(f"    ✗ Error: {e}")
        results.append(("Custom Headers/Footers", False))

    # Test 4: Full document conversion with page numbers
    print("\n  → Testing PDF conversion with page numbers...")
    try:
        # Create test markdown with multiple pages
        pagenums_dir = TEST_DIR / "page_numbers"
        pagenums_dir.mkdir(exist_ok=True)

        test_md = pagenums_dir / "test_with_pages.md"
        test_content = """---
title: "Test Document with Page Numbers"
author: "E-Book Maker Test Suite"
---

# Chapter 1: Introduction

This is a test document to verify page numbering functionality.

Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.

\\newpage

# Chapter 2: Content

Second page with more content to test page numbering.

Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.

\\newpage

# Chapter 3: Conclusion

Third page to ensure page numbers appear on all pages.

Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.
"""
        test_md.write_text(test_content, encoding='utf-8')

        # Create converter with page numbering
        page_config = {
            'pdf': {
                'enabled': True,
                'position': 'footer-center',
                'style': 'arabic'
            }
        }

        converter = EBookConverter(page_numbering_config=page_config)

        # Convert to PDF
        pdf_output = converter.convert_to_pdf(test_md, pagenums_dir)

        if pdf_output and pdf_output.exists():
            size = pdf_output.stat().st_size
            print(f"    ✓ PDF created with page numbers: {pdf_output.name} ({size:,} bytes)")
            results.append(("PDF with Page Numbers", True))
        else:
            print(f"    ⚠ PDF conversion completed but file verification failed")
            results.append(("PDF with Page Numbers", False))

    except Exception as e:
        print(f"    ✗ Error: {e}")
        import traceback
        traceback.print_exc()
        results.append(("PDF with Page Numbers", False))

    # Summary
    passed = sum(1 for _, success in results if success)
    total = len(results)
    print(f"\n  PAGE NUMBERING: {passed}/{total} tests passed")

    return results


# Run all tests
if __name__ == "__main__":
    all_results = {}

    # Test 1: Watermarking
    watermark_results = test_watermarking()
    all_results['watermarking'] = watermark_results

    # Test 2: Text Normalization
    normalize_results = test_text_normalization()
    all_results['text_normalization'] = normalize_results

    # Test 3: Page Numbering
    pagenum_results = test_page_numbering()
    all_results['page_numbering'] = pagenum_results

    # Final summary
    print("\n" + "=" * 80)
    print("📊 FINAL SUMMARY")
    print("=" * 80)

    total_passed = 0
    total_tests = 0

    for category, results in all_results.items():
        if isinstance(results[0], tuple) and len(results[0]) >= 2:
            passed = sum(1 for r in results if r[1])
            total = len(results)
            total_passed += passed
            total_tests += total

            status = "✓" if passed == total else "⚠"
            print(f"{status} {category.replace('_', ' ').title()}: {passed}/{total}")

    percentage = (total_passed / total_tests * 100) if total_tests > 0 else 0
    print(f"\n  OVERALL: {total_passed}/{total_tests} ({percentage:.1f}%)")

    print(f"\n📁 All test outputs saved to:")
    print(f"   {TEST_DIR}")
    print("=" * 80)
